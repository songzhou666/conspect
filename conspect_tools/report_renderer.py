"""
vision_tools 报告渲染引擎模块
将报告设计渲染为多种格式输出（Markdown、HTML、PDF、Word）

核心设计原则：
  - 支持多种输出格式
  - 内联所有 CSS 样式，支持离线阅读
  - 支持目录导航和打印优化
"""
from typing import Dict, List, Optional, Union
from pathlib import Path


class ReportRenderer:
    """
    报告渲染引擎
    将报告设计渲染为多种格式输出
    
    使用示例:
        renderer = ReportRenderer(theme_name="ocean")
        markdown = renderer.render_markdown(report_design)
        html = renderer.render_html(report_design)
    """
    
    def __init__(self, theme_name: str = "ocean"):
        """
        初始化报告渲染引擎
        
        参数:
            theme_name: 配色主题名称
        """
        # 延迟导入，避免循环依赖
        from .render_engine import ColorTheme
        self.theme = ColorTheme.get(theme_name)
    
    def render_markdown(self, report_design: Dict) -> str:
        """
        渲染为 Markdown 格式
        
        参数:
            report_design: 报告设计字典
            
        返回:
            Markdown 字符串
        """
        sections = [
            self._render_header(report_design),
            self._render_executive_summary(report_design),
            self._render_data_overview(report_design),
            self._render_key_findings(report_design),
            self._render_advanced_analysis(report_design),
            self._render_recommendations(report_design),
            self._render_appendix(report_design)
        ]
        return "\n\n---\n\n".join(sections)
    
    def render_html(self, report_design: Dict) -> str:
        """
        渲染为 HTML 格式（独立文件，可离线阅读）
        
        特点:
        - 内联所有 CSS 样式
        - 支持目录导航
        - 支持打印优化
        """
        md_content = self.render_markdown(report_design)
        html = self._md_to_html(md_content)
        return self._wrap_html_template(html, report_design.get("title", "数据分析报告"))
    
    def render_pdf(self, report_design: Dict) -> bytes:
        """
        渲染为 PDF 格式
        
        实现: 通过 Playwright 无头浏览器打开 HTML，截图输出 PDF
        
        注意: 需要安装 playwright: pip install playwright && playwright install chromium
        """
        html = self.render_html(report_design)
        
        try:
            from playwright.sync_api import sync_playwright
            import io
            
            with sync_playwright() as p:
                browser = p.chromium.launch()
                page = browser.new_page()
                page.set_content(html)
                pdf = page.pdf(
                    format='A4',
                    margin={
                        'top': '20mm',
                        'right': '20mm',
                        'bottom': '20mm',
                        'left': '20mm'
                    },
                    print_background=True
                )
                browser.close()
                return pdf
        except ImportError:
            # 如果未安装 playwright，返回 HTML 格式的字节
            return html.encode('utf-8')
        except Exception as e:
            # 如果 PDF 生成失败，返回 HTML 格式的字节
            return html.encode('utf-8')
    
    def render_word(self, report_design: Dict) -> bytes:
        """
        渲染为 Word 格式 (.docx)
        
        实现: 使用 python-docx 库生成
        
        注意: 需要安装 python-docx: pip install python-docx
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            # 如果未安装 python-docx，返回 Markdown 格式的字节
            markdown = self.render_markdown(report_design)
            return markdown.encode('utf-8')
        
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)
        
        # 添加标题
        title = doc.add_heading(report_design.get('title', '数据分析报告'), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加元信息
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"生成时间: {report_design.get('generated_at', '')}\n")
        meta_para.add_run(f"数据来源: {report_design.get('data_source', '')}\n")
        meta_para.add_run(f"数据范围: {report_design.get('data_range', '')}\n")
        
        doc.add_paragraph()  # 空行
        
        # 添加执行摘要
        executive_summary = report_design.get('executive_summary', '')
        if executive_summary:
            doc.add_heading('执行摘要', level=1)
            doc.add_paragraph(executive_summary)
        
        # 添加关键发现
        key_findings = report_design.get('key_findings', [])
        if key_findings:
            doc.add_heading('关键发现', level=1)
            for finding in key_findings:
                if hasattr(finding, 'title'):
                    doc.add_heading(finding.title, level=2)
                    doc.add_paragraph(finding.description)
                    doc.add_paragraph(f"数据依据: {finding.evidence}")
                else:
                    doc.add_paragraph(str(finding))
        
        # 添加行动建议
        recommendations = report_design.get('recommendations', [])
        if recommendations:
            doc.add_heading('行动建议', level=1)
            for i, rec in enumerate(recommendations, 1):
                if hasattr(rec, 'title'):
                    doc.add_paragraph(f"{i}. {rec.title} [{rec.priority}]")
                    doc.add_paragraph(rec.description)
                else:
                    doc.add_paragraph(f"{i}. {str(rec)}")
        
        # 保存到字节流
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _render_header(self, design: Dict) -> str:
        """
        渲染报告头部
        
        参数:
            design: 报告设计字典
            
        返回:
            Markdown 格式的头部
        """
        title = design.get("title", "数据分析报告")
        generated_at = design.get("generated_at", "")
        data_source = design.get("data_source", "")
        data_range = design.get("data_range", "")
        
        return f"""# {title}

> **生成时间**: {generated_at}  
> **数据来源**: {data_source}  
> **数据范围**: {data_range}
"""
    
    def _render_executive_summary(self, design: Dict) -> str:
        """
        渲染执行摘要
        
        参数:
            design: 报告设计字典
            
        返回:
            Markdown 格式的执行摘要
        """
        summary = design.get("executive_summary", "")
        kpis = design.get("kpis", [])
        
        kpi_text = "\n".join([
            f"| {kpi.get('label', '')} | {kpi.get('value', '')} | {kpi.get('change', '')} |"
            for kpi in kpis
        ])
        
        return f"""## 执行摘要

{summary}

### 关键指标

| 指标 | 数值 | 变化 |
|------|------|------|
{kpi_text}
"""
    
    def _render_data_overview(self, design: Dict) -> str:
        """
        渲染数据概览
        
        参数:
            design: 报告设计字典
            
        返回:
            Markdown 格式的数据概览
        """
        data_overview = design.get("data_overview", {})
        
        if not data_overview:
            return ""
        
        rows = data_overview.get("rows", 0)
        columns = data_overview.get("columns", 0)
        data_range = data_overview.get("data_range", "")
        
        return f"""## 数据概览

- **数据量**: {rows} 行 × {columns} 列
- **数据范围**: {data_range}
"""
    
    def _render_key_findings(self, design: Dict) -> str:
        """
        渲染关键发现
        
        参数:
            design: 报告设计字典
            
        返回:
            Markdown 格式的关键发现
        """
        findings = design.get("key_findings", [])
        
        if not findings:
            return ""
        
        sections = []
        for finding in findings:
            if hasattr(finding, 'to_markdown'):
                sections.append(finding.to_markdown())
            else:
                sections.append(str(finding))
        
        return "## 关键发现\n\n" + "\n\n".join(sections)
    
    def _render_advanced_analysis(self, design: Dict) -> str:
        """
        渲染高级分析
        
        参数:
            design: 报告设计字典
            
        返回:
            Markdown 格式的高级分析
        """
        advanced_analysis = design.get("advanced_analysis", {})
        
        if not advanced_analysis:
            return ""
        
        sections = []
        
        # 信息熵分析
        entropy = advanced_analysis.get("entropy", {})
        if entropy:
            sections.append("### 信息熵分析\n\n" + str(entropy))
        
        # 集中度分析
        concentration = advanced_analysis.get("concentration", {})
        if concentration:
            sections.append("### 集中度分析\n\n" + str(concentration))
        
        # 趋势预测
        trend = advanced_analysis.get("trend", {})
        if trend:
            sections.append("### 趋势预测\n\n" + str(trend))
        
        if not sections:
            return ""
        
        return "## 高级分析\n\n" + "\n\n".join(sections)
    
    def _render_recommendations(self, design: Dict) -> str:
        """
        渲染行动建议
        
        参数:
            design: 报告设计字典
            
        返回:
            Markdown 格式的行动建议
        """
        recommendations = design.get("recommendations", [])
        
        if not recommendations:
            return ""
        
        # 按优先级排序
        priority_order = {"high": 0, "medium": 1, "low": 2}
        recommendations.sort(key=lambda x: priority_order.get(
            getattr(x, 'priority', 'low') if hasattr(x, 'priority') else x.get('priority', 'low'), 
            3
        ))
        
        sections = []
        for rec in recommendations:
            if hasattr(rec, 'to_markdown'):
                sections.append(rec.to_markdown())
            else:
                sections.append(str(rec))
        
        return "## 行动建议\n\n" + "\n\n".join(sections)
    
    def _render_appendix(self, design: Dict) -> str:
        """
        渲染附录
        
        参数:
            design: 报告设计字典
            
        返回:
            Markdown 格式的附录
        """
        appendix = design.get("appendix", {})
        
        if not appendix:
            return ""
        
        return f"""## 附录

### 分析方法

{appendix.get('analysis_method', '高级统计分析')}

### 数据源

{appendix.get('data_source', '')}

### 数据范围

{appendix.get('data_range', '')}
"""
    
    def _md_to_html(self, md_content: str) -> str:
        """
        Markdown 转 HTML
        
        参数:
            md_content: Markdown 内容
            
        返回:
            HTML 字符串
            
        注意: 这是一个简化实现，实际使用时可以安装 markdown 库
        """
        # 简化实现：基本 Markdown 转 HTML
        # 实际使用时建议安装 markdown 库: pip install markdown
        try:
            import markdown
            return markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
        except ImportError:
            # 简化实现
            html = md_content
            # 标题
            html = html.replace("# ", "<h1>").replace("\n", "</h1>\n", 1)
            html = html.replace("## ", "<h2>").replace("\n", "</h2>\n", 1)
            html = html.replace("### ", "<h3>").replace("\n", "</h3>\n", 1)
            # 粗体
            html = html.replace("**", "<strong>", 1)
            html = html.replace("**", "</strong>", 1)
            # 表格（简化）
            html = html.replace("|", "</td><td>")
            # 引用
            html = html.replace("> ", "<blockquote>")
            return html
    
    def _wrap_html_template(self, content: str, title: str) -> str:
        """
        包装为完整 HTML 页面
        
        参数:
            content: HTML 内容
            title: 页面标题
            
        返回:
            完整的 HTML 页面
        """
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        /* 报告专用样式 */
        body {{
            font-family: 'PingFang SC', 'Microsoft YaHei', -apple-system, sans-serif;
            line-height: 1.8;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
        }}
        h1 {{ 
            color: #2B5F8A; 
            border-bottom: 2px solid #2B5F8A; 
            padding-bottom: 10px; 
        }}
        h2 {{ 
            color: #2B5F8A; 
            margin-top: 40px; 
            border-bottom: 1px solid #E2E8F0;
            padding-bottom: 8px;
        }}
        h3 {{ 
            color: #4A9BD9; 
            margin-top: 30px; 
        }}
        table {{ 
            border-collapse: collapse; 
            width: 100%; 
            margin: 20px 0; 
        }}
        th, td {{ 
            border: 1px solid #ddd; 
            padding: 12px; 
            text-align: left; 
        }}
        th {{ 
            background-color: #f5f5f5; 
            font-weight: 600;
        }}
        tr:nth-child(even) {{
            background-color: #f9f9f9;
        }}
        blockquote {{ 
            border-left: 4px solid #2B5F8A; 
            padding: 12px 16px; 
            color: #666; 
            background-color: #f8f9fa;
            margin: 16px 0;
        }}
        .severity-info {{ color: #4A9BD9; }}
        .severity-warning {{ color: #E8856B; }}
        .severity-critical {{ color: #D32F2F; }}
        ul, ol {{
            padding-left: 24px;
        }}
        li {{
            margin: 8px 0;
        }}
        @media print {{
            body {{ 
                padding: 0; 
                max-width: 100%;
            }}
            .no-print {{ 
                display: none; 
            }}
            h1, h2, h3 {{
                page-break-after: avoid;
            }}
            table {{
                page-break-inside: avoid;
            }}
        }}
    </style>
</head>
<body>
    {content}
</body>
</html>"""
